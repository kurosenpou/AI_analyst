"""
File parsing service
Handles CSV, PDF, and Excel file parsing with data extraction
"""

import pandas as pd
import PyPDF2
import os
from typing import Dict, Any, List, Union
import json
import docx
from docx import Document
import chardet
import re
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class FileParser:
    """Handles parsing of uploaded files into structured data"""
    
    def __init__(self):
        self.supported_formats = ['.csv', '.pdf', '.xlsx', '.xls', '.json', '.txt', '.docx']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.encoding_fallbacks = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'gbk']
    
    async def parse_file(self, file_path: str, file_extension: str) -> Union[List[Dict], Dict[str, Any]]:
        """
        Parse file based on its extension
        Returns structured data ready for AI analysis
        """
        
        # Validate file size
        if os.path.getsize(file_path) > self.max_file_size:
            raise ValueError(f"File size exceeds maximum limit of {self.max_file_size / (1024*1024):.1f}MB")
        
        try:
            if file_extension == '.csv':
                return await self._parse_csv(file_path)
            elif file_extension == '.pdf':
                return await self._parse_pdf(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return await self._parse_excel(file_path)
            elif file_extension == '.json':
                return await self._parse_json(file_path)
            elif file_extension == '.txt':
                return await self._parse_txt(file_path)
            elif file_extension == '.docx':
                return await self._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}")
            raise Exception(f"Failed to parse {file_extension} file: {str(e)}")
    
    async def _parse_csv(self, file_path: str) -> List[Dict]:
        """Parse CSV file into list of dictionaries"""
        try:
            df = pd.read_csv(file_path)
            
            # Clean data: remove empty rows, handle NaN values
            df = df.dropna(how='all')  # Remove completely empty rows
            df = df.fillna('')  # Fill NaN with empty string
            
            # Convert to list of dictionaries
            data = df.to_dict('records')
            
            return data
            
        except Exception as e:
            raise Exception(f"Error parsing CSV file: {str(e)}")
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file and extract text content"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_content = ""
                page_contents = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content += page_text + "\n"
                    page_contents.append({
                        "page_number": page_num + 1,
                        "content": page_text.strip()
                    })
                
                return {
                    "total_pages": len(pdf_reader.pages),
                    "full_text": text_content.strip(),
                    "pages": page_contents,
                    "metadata": {
                        "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                        "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                        "subject": pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else ''
                    }
                }
                
        except Exception as e:
            raise Exception(f"Error parsing PDF file: {str(e)}")
    
    async def _parse_excel(self, file_path: str) -> List[Dict]:
        """Parse Excel file into list of dictionaries"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            if len(excel_data) == 1:
                # Single sheet - return as list of records
                sheet_name = list(excel_data.keys())[0]
                df = excel_data[sheet_name]
                df = df.dropna(how='all').fillna('')
                return df.to_dict('records')
            else:
                # Multiple sheets - return structured data
                result = {}
                for sheet_name, df in excel_data.items():
                    df = df.dropna(how='all').fillna('')
                    result[sheet_name] = df.to_dict('records')
                return result
                
        except Exception as e:
            raise Exception(f"Error parsing Excel file: {str(e)}")
    
    def create_data_preview(self, parsed_data: Union[List[Dict], Dict[str, Any]]) -> Dict[str, Any]:
        """
        Create a preview of parsed data for frontend display
        Limits data size to avoid overwhelming the UI
        """
        
        if isinstance(parsed_data, list):
            # CSV/Excel/JSON array data
            preview = {
                "type": "tabular",
                "total_rows": len(parsed_data),
                "columns": list(parsed_data[0].keys()) if parsed_data else [],
                "sample_rows": parsed_data[:5],  # First 5 rows
                "data_types": self._analyze_data_types(parsed_data[:100])  # Analyze first 100 rows
            }
            
        elif isinstance(parsed_data, dict):
            if "content_type" in parsed_data:
                content_type = parsed_data["content_type"]
                
                if content_type == "text":
                    # TXT file data
                    preview = {
                        "type": "text",
                        "line_count": parsed_data.get("line_count", 0),
                        "word_count": parsed_data.get("word_count", 0),
                        "character_count": parsed_data.get("character_count", 0),
                        "text_preview": parsed_data["full_text"][:500] + "..." if len(parsed_data["full_text"]) > 500 else parsed_data["full_text"],
                        "structured_data": parsed_data.get("structured_data", {}),
                        "metadata": parsed_data.get("metadata", {})
                    }
                    
                elif content_type == "document":
                    # DOCX file data
                    preview = {
                        "type": "document",
                        "word_count": parsed_data.get("word_count", 0),
                        "paragraph_count": parsed_data.get("paragraph_count", 0),
                        "table_count": parsed_data.get("table_count", 0),
                        "text_preview": parsed_data["full_text"][:500] + "..." if len(parsed_data["full_text"]) > 500 else parsed_data["full_text"],
                        "tables_preview": [table[:3] for table in parsed_data.get("tables", [])[:2]],  # First 2 tables, 3 rows each
                        "metadata": parsed_data.get("metadata", {})
                    }
                else:
                    preview = {"type": "unknown", "data": str(parsed_data)[:200]}
                    
            elif "full_text" in parsed_data and "total_pages" in parsed_data:
                # PDF data
                preview = {
                    "type": "pdf",
                    "total_pages": parsed_data.get("total_pages", 0),
                    "text_preview": parsed_data["full_text"][:500] + "..." if len(parsed_data["full_text"]) > 500 else parsed_data["full_text"],
                    "metadata": parsed_data.get("metadata", {}),
                    "word_count": len(parsed_data["full_text"].split())
                }
                
            elif all(isinstance(v, list) for v in parsed_data.values()):
                # Multi-sheet Excel or flattened JSON
                preview = {
                    "type": "multi_sheet",
                    "sheets": {
                        sheet_name: {
                            "rows": len(sheet_data),
                            "columns": list(sheet_data[0].keys()) if sheet_data and isinstance(sheet_data[0], dict) else [],
                            "sample": sheet_data[:3] if isinstance(sheet_data, list) else []
                        }
                        for sheet_name, sheet_data in parsed_data.items()
                    }
                }
                
            else:
                # Flattened JSON object or other structured data
                preview = {
                    "type": "structured",
                    "keys": list(parsed_data.keys())[:20],  # First 20 keys
                    "sample_data": {k: v for k, v in list(parsed_data.items())[:10]},  # First 10 key-value pairs
                    "total_keys": len(parsed_data)
                }
                
        else:
            preview = {"type": "unknown", "data": str(parsed_data)[:200]}
        
        return preview
    
    def _analyze_data_types(self, data_sample: List[Dict]) -> Dict[str, str]:
        """Analyze data types of columns for better AI understanding"""
        if not data_sample:
            return {}
        
        columns = data_sample[0].keys()
        type_analysis = {}
        
        for col in columns:
            values = [row.get(col, '') for row in data_sample if row.get(col, '') != '']
            
            if not values:
                type_analysis[col] = "empty"
                continue
            
            # Try to determine type
            numeric_count = sum(1 for v in values if str(v).replace('.', '').replace('-', '').isdigit())
            date_count = sum(1 for v in values if self._is_date_like(str(v)))
            
            if numeric_count / len(values) > 0.8:
                type_analysis[col] = "numeric"
            elif date_count / len(values) > 0.8:
                type_analysis[col] = "date"
            else:
                type_analysis[col] = "text"
        
        return type_analysis
    
    def _is_date_like(self, value: str) -> bool:
        """Simple date detection"""
        date_indicators = ['/', '-', ' ', '年', '月', '日']
        return any(indicator in value for indicator in date_indicators) and any(c.isdigit() for c in value)
    
    async def _parse_json(self, file_path: str) -> Union[List[Dict], Dict[str, Any]]:
        """Parse JSON file into structured data"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding) as file:
                data = json.load(file)
            
            # Validate and structure JSON data
            if isinstance(data, list):
                # Array of objects - treat as tabular data
                return self._normalize_json_array(data)
            elif isinstance(data, dict):
                # Single object or nested structure
                return self._flatten_json_object(data)
            else:
                # Primitive type - wrap in structure
                return {"data": data, "type": type(data).__name__}
                
        except json.JSONDecodeError as e:
            raise Exception(f"Invalid JSON format: {str(e)}")
        except Exception as e:
            raise Exception(f"Error parsing JSON file: {str(e)}")
    
    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Parse text file and extract structured information"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            # Analyze text structure
            lines = content.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            
            # Try to detect structured data patterns
            structured_data = self._analyze_text_structure(content, non_empty_lines)
            
            return {
                "content_type": "text",
                "full_text": content,
                "line_count": len(lines),
                "word_count": len(content.split()),
                "character_count": len(content),
                "structured_data": structured_data,
                "metadata": {
                    "encoding": encoding,
                    "has_tables": self._detect_tables_in_text(non_empty_lines),
                    "has_lists": self._detect_lists_in_text(non_empty_lines),
                    "sections": self._extract_sections(non_empty_lines)
                }
            }
            
        except Exception as e:
            raise Exception(f"Error parsing text file: {str(e)}")
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file and extract text and structure"""
        try:
            doc = Document(file_path)
            
            # Extract text content
            full_text = ""
            paragraphs = []
            tables_data = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal"
                    })
                    full_text += para.text + "\n"
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_data.append(row_data)
                if table_data:
                    tables_data.append(table_data)
            
            return {
                "content_type": "document",
                "full_text": full_text.strip(),
                "paragraphs": paragraphs,
                "tables": tables_data,
                "word_count": len(full_text.split()),
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_data),
                "metadata": {
                    "core_properties": self._extract_docx_properties(doc),
                    "styles_used": list(set(p["style"] for p in paragraphs))
                }
            }
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")
    
    def _detect_encoding(self, file_path: str) -> str:
        """Detect file encoding with fallback options"""
        try:
            with open(file_path, 'rb') as file:
                raw_data = file.read(10000)  # Read first 10KB for detection
                result = chardet.detect(raw_data)
                if result['encoding'] and result['confidence'] > 0.7:
                    return result['encoding']
        except:
            pass
        
        # Try fallback encodings
        for encoding in self.encoding_fallbacks:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    file.read(1000)  # Test read
                return encoding
            except:
                continue
        
        return 'utf-8'  # Final fallback
    
    def _normalize_json_array(self, data: List) -> List[Dict]:
        """Normalize JSON array to consistent structure"""
        if not data:
            return []
        
        # If all items are dictionaries, return as is
        if all(isinstance(item, dict) for item in data):
            return data
        
        # Convert non-dict items to dict format
        normalized = []
        for i, item in enumerate(data):
            if isinstance(item, dict):
                normalized.append(item)
            else:
                normalized.append({"index": i, "value": item, "type": type(item).__name__})
        
        return normalized
    
    def _flatten_json_object(self, obj: Dict, prefix: str = "") -> Dict[str, Any]:
        """Flatten nested JSON object"""
        flattened = {}
        
        for key, value in obj.items():
            new_key = f"{prefix}.{key}" if prefix else key
            
            if isinstance(value, dict):
                flattened.update(self._flatten_json_object(value, new_key))
            elif isinstance(value, list):
                if value and isinstance(value[0], dict):
                    # Array of objects - keep as nested structure
                    flattened[new_key] = value
                else:
                    # Array of primitives - flatten
                    for i, item in enumerate(value):
                        flattened[f"{new_key}[{i}]"] = item
            else:
                flattened[new_key] = value
        
        return flattened
    
    def _analyze_text_structure(self, content: str, lines: List[str]) -> Dict[str, Any]:
        """Analyze text structure to extract meaningful data"""
        structure = {
            "key_value_pairs": [],
            "numbered_lists": [],
            "bullet_points": [],
            "sections": []
        }
        
        # Extract key-value pairs (key: value, key = value, etc.)
        kv_pattern = r'^([^:=]+)[:=]\s*(.+)$'
        for line in lines:
            match = re.match(kv_pattern, line.strip())
            if match:
                structure["key_value_pairs"].append({
                    "key": match.group(1).strip(),
                    "value": match.group(2).strip()
                })
        
        # Extract numbered lists
        num_pattern = r'^\d+[.)\s]+(.+)$'
        for line in lines:
            match = re.match(num_pattern, line.strip())
            if match:
                structure["numbered_lists"].append(match.group(1).strip())
        
        # Extract bullet points
        bullet_pattern = r'^[-*•]\s+(.+)$'
        for line in lines:
            match = re.match(bullet_pattern, line.strip())
            if match:
                structure["bullet_points"].append(match.group(1).strip())
        
        return structure
    
    def _detect_tables_in_text(self, lines: List[str]) -> bool:
        """Detect if text contains table-like structures"""
        table_indicators = ['|', '\t', '  ']
        table_lines = 0
        
        for line in lines:
            if any(indicator in line for indicator in table_indicators):
                # Check if line has multiple separated values
                parts = re.split(r'[|\t]|\s{2,}', line)
                if len(parts) > 2:
                    table_lines += 1
        
        return table_lines > 2
    
    def _detect_lists_in_text(self, lines: List[str]) -> bool:
        """Detect if text contains list structures"""
        list_patterns = [r'^\d+[.)\s]', r'^[-*•]\s', r'^[a-zA-Z][.)\s]']
        list_lines = 0
        
        for line in lines:
            if any(re.match(pattern, line.strip()) for pattern in list_patterns):
                list_lines += 1
        
        return list_lines > 1
    
    def _extract_sections(self, lines: List[str]) -> List[str]:
        """Extract section headers from text"""
        sections = []
        
        for line in lines:
            # Look for lines that might be headers (all caps, short, etc.)
            if (len(line) < 100 and 
                (line.isupper() or 
                 line.endswith(':') or 
                 re.match(r'^\d+\.\s+[A-Z]', line) or
                 re.match(r'^[A-Z][^.!?]*$', line))):
                sections.append(line.strip())
        
        return sections[:10]  # Limit to first 10 potential sections
    
    def _extract_docx_properties(self, doc: Document) -> Dict[str, Any]:
        """Extract document properties from DOCX"""
        try:
            props = doc.core_properties
            return {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "keywords": props.keywords or ""
            }
        except:
            return {}
